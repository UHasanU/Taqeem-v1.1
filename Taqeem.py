from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from typing import Optional


class Taqeem:
    def __init__(self):
        self.ws = None
        self.doc = Document()
        self.instructors = set()
        self.instructors_column = ''
        self.sorted_columns = ()

    def add_resource(self, workbook_name, worksheet_name):
        self.ws = load_workbook(workbook_name)[worksheet_name]

    def set_instructors_column(self, instructors_column):
        self.instructors_column = instructors_column

        for instructor in self.ws[instructors_column]:
            if instructor.row == 1 or instructor.value == "مو موجود":
                continue
            self.instructors.add(instructor.value)
        self.instructors = sorted(self.instructors)

    def set_columns_to_display(self, *columns):
        self.sorted_columns = columns

    def add_title_page(self,
                       title: str,
                       font: Optional[str] = None,
                       font_size: Optional[float] = 24,
                       margins: Optional[float] = 0.5,
                       picture: Optional[str] = None
                       ):
        if font is None:
            if "a" <= title[0] <= "z" or "A" <= title[0] <= "Z":
                font = "Times New Roman"
            else:
                font = "DecoType Naskh"

        section = self.doc.sections[0]
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

        title_para = self.doc.add_paragraph()
        Taqeem.set_paragraph_text(title_para, title, font, font_size, WD_PARAGRAPH_ALIGNMENT.CENTER, True)

        if picture is not None:
            pass  # Later

        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def add_taqeemat(self,
                     table_style: Optional[str] = "Light List Accent 2",
                     font: Optional[str] = "Times New Roman",
                     font_size: Optional[float] = 16
                     ):
        for instructor_name in self.instructors:
            inst_name = self.doc.add_heading(level=1)
            Taqeem.set_paragraph_text(inst_name, instructor_name.title(), font, font_size,
                                      WD_PARAGRAPH_ALIGNMENT.CENTER, True)

            instructor_rows = [instructor.row for instructor in self.ws[self.instructors_column] if
                               instructor.value == instructor_name]

            for row_number in instructor_rows:
                table = self.doc.add_table(rows=8, cols=2)
                table.style = table_style
                table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title_cell = table.cell(0, 0).merge(table.cell(0, 1))
                Taqeem.set_paragraph_text(title_cell.paragraphs[0], instructor_name.title(), font,
                                          font_size, WD_PARAGRAPH_ALIGNMENT.CENTER, bold=True)

                for count, row in enumerate(table.rows[:-1]):
                    row.cells[0].width = Inches(5.25)
                    row.cells[1].width = Inches(1.86)

                    right_cell_value = str(self.ws[f"{self.sorted_columns[count]}1"].value)
                    left_cell_value = Taqeem.fix_imlaa(str(self.ws[f"{self.sorted_columns[count]}{row_number}"].value))
                    Taqeem.set_paragraph_text(table.cell(count + 1, 1).paragraphs[0], right_cell_value, font, font_size,
                                              WD_PARAGRAPH_ALIGNMENT.RIGHT)
                    Taqeem.set_paragraph_text(table.cell(count + 1, 0).paragraphs[0], left_cell_value, font, font_size,
                                              WD_PARAGRAPH_ALIGNMENT.RIGHT)

                self.doc.add_paragraph().add_run()
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self, filename):
        self.doc.save(filename)

    @staticmethod
    def fix_imlaa(text: str):
        misspelled_words = {"مره", "مرا", "سهله", "اجوبه", "أجوبه", "فابده", "ساعه", "صراحه", "ماده", "دقيقه", "نهايه",
                            "حصه", "محاضره"}
        word_list = text.split()
        for index, word in enumerate(word_list):
            if word in misspelled_words or (word.startswith("ال") and word[2:] in misspelled_words):
                word_list[index] = word[:-1] + "ة"
        return " ".join(word_list)

    @staticmethod
    def set_paragraph_text(paragraph, text, font, font_size, alignment, bold=False):
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        run.font.name = font
        run.font.size = Pt(font_size)
        run.font.bold = bold
